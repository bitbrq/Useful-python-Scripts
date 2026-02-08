import os
import sys
import subprocess
from pathlib import Path

def install_package(package_name):
    """Install a package using pip"""
    try:
        print(f"Installing {package_name}...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", package_name])
        return True
    except subprocess.CalledProcessError:
        print(f"Failed to install {package_name}")
        return False

def check_and_install_packages():
    """Check and install required packages"""
    required_packages = ['PyPDF2', 'python-docx']
    
    for package in required_packages:
        try:
            if package == 'PyPDF2':
                import PyPDF2
            elif package == 'python-docx':
                import docx
            print(f"{package} is already installed")
        except ImportError:
            install_package(package)

def convert_pdf_to_word_simple(pdf_path, docx_path):
    """
    Convert PDF to Word using PyPDF2 and python-docx (simple text extraction)
    This preserves basic text structure but not complex formatting
    """
    try:
        print(f"Converting {os.path.basename(pdf_path)} to Word format...")
        
        # Import inside function to handle import errors gracefully
        from PyPDF2 import PdfReader
        from docx import Document
        from docx.shared import Pt
        
        # Open the PDF file
        pdf_reader = PdfReader(pdf_path)
        
        # Create a new Word document
        doc = Document()
        
        # Process each page
        for page_num, page in enumerate(pdf_reader.pages):
            # Extract text from the page
            text = page.extract_text()
            
            # Skip empty pages
            if not text.strip():
                continue
                
            # Add a heading for the page (except first page)
            if page_num > 0:
                doc.add_page_break()
                doc.add_heading(f"Page {page_num + 1}", level=2)
            
            # Split text into paragraphs and add to document
            paragraphs = text.split('\n')
            for para in paragraphs:
                if para.strip():  # Skip empty paragraphs
                    p = doc.add_paragraph(para.strip())
                    # Set font size
                    for run in p.runs:
                        run.font.size = Pt(11)
        
        # Save the document
        doc.save(docx_path)
        
        print(f"✓ Successfully converted to: {os.path.basename(docx_path)}")
        print(f"  Location: {docx_path}")
        return True
        
    except Exception as e:
        print(f"✗ Error during conversion: {str(e)}")
        return False

def find_pdf_files():
    """Find all PDF files in the current directory"""
    pdf_files = []
    current_dir = os.getcwd()
    
    for file in os.listdir(current_dir):
        if file.lower().endswith('.pdf'):
            pdf_files.append(os.path.join(current_dir, file))
    
    return pdf_files

def main():
    """Main function to convert PDF to Word"""
    print("=" * 60)
    print("PDF to Word Converter")
    print("=" * 60)
    
    # Check and install required packages
    check_and_install_packages()
    
    # Find PDF files
    pdf_files = find_pdf_files()
    
    if not pdf_files:
        print("No PDF files found in the current directory.")
        print("Please place your PDF file in the same directory as this script.")
        return
    
    print(f"\nFound {len(pdf_files)} PDF file(s):")
    for i, pdf_file in enumerate(pdf_files, 1):
        print(f"  {i}. {os.path.basename(pdf_file)}")
    
    # Process each PDF file
    for pdf_path in pdf_files:
        pdf_name = os.path.basename(pdf_path)
        docx_name = os.path.splitext(pdf_name)[0] + ".docx"
        docx_path = os.path.join(os.path.dirname(pdf_path), docx_name)
        
        print(f"\n{'='*60}")
        print(f"Processing: {pdf_name}")
        print(f"{'='*60}")
        
        # Ask for confirmation if file exists
        if os.path.exists(docx_path):
            response = input(f"\n{os.path.basename(docx_path)} already exists. Overwrite? (y/n): ")
            if response.lower() != 'y':
                print("Skipping this file...")
                continue
        
        # Convert the file
        success = convert_pdf_to_word_simple(pdf_path, docx_path)
        
        if success:
            print(f"\n✓ Conversion completed successfully!")
        else:
            print(f"\n✗ Failed to convert {pdf_name}")
    
    print("\n" + "=" * 60)
    print("Conversion process completed!")
    print("=" * 60)
    
    # Wait for user input before closing
    input("\nPress Enter to exit...")

if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n\nProcess interrupted by user.")
    except Exception as e:
        print(f"\nAn unexpected error occurred: {str(e)}")
        input("Press Enter to exit...")
        